"""
Napkin — Project & Feedback Routes
CRUD for projects and feedback ingestion.
"""

from typing import Annotated
from uuid import UUID

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, status

from app.api.deps.auth import get_current_user
from app.db.client import get_supabase_admin
from app.schemas.api import (
    FeedbackItemResponse,
    FeedbackListResponse,
    FeedbackPaste,
    FeedbackUploadResponse,
    FileParseResponse,
    ProjectCreate,
    ProjectUpdate,
)

MAX_UPLOAD_SIZE = 10 * 1024 * 1024  # 10 MB

# ============================================================
# PROJECTS
# ============================================================

projects_router = APIRouter(prefix="/projects", tags=["projects"])


def _verify_project_access(db, project_id: UUID, user: dict) -> dict:
    """Verify the user's org owns the project. Returns the project."""
    result = db.table("projects").select("*").eq("id", str(project_id)).single().execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Project not found")
    if result.data.get("org_id") != user.get("org_id"):
        raise HTTPException(status_code=403, detail="Access denied")
    return result.data


@projects_router.post("", response_model=dict, status_code=status.HTTP_201_CREATED)
async def create_project(
    body: ProjectCreate,
    user: Annotated[dict, Depends(get_current_user)],
):
    """Create a new project within the user's organization."""
    if not user.get("org_id"):
        raise HTTPException(status_code=400, detail="Must belong to an organization")

    db = get_supabase_admin()
    data = {
        "org_id": user["org_id"],
        "name": body.name,
        "description": body.description,
        "repo_url": body.repo_url,
        "repo_provider": body.repo_provider,
        "created_by": user["id"],
    }

    result = db.table("projects").insert(data).execute()
    return result.data[0] if result.data else {}


@projects_router.get("", response_model=list[dict])
async def list_projects(
    user: Annotated[dict, Depends(get_current_user)],
):
    """List all projects in the user's organization."""
    if not user.get("org_id"):
        return []

    db = get_supabase_admin()
    result = (
        db.table("projects")
        .select("*")
        .eq("org_id", user["org_id"])
        .order("created_at", desc=True)
        .execute()
    )
    return result.data


@projects_router.get("/{project_id}", response_model=dict)
async def get_project(
    project_id: UUID,
    user: Annotated[dict, Depends(get_current_user)],
):
    """Get project details."""
    db = get_supabase_admin()
    return _verify_project_access(db, project_id, user)


@projects_router.patch("/{project_id}", response_model=dict)
async def update_project(
    project_id: UUID,
    body: ProjectUpdate,
    user: Annotated[dict, Depends(get_current_user)],
):
    """Update project settings."""
    db = get_supabase_admin()
    _verify_project_access(db, project_id, user)

    update_data = body.model_dump(exclude_none=True)
    if not update_data:
        raise HTTPException(status_code=400, detail="No fields to update")

    result = db.table("projects").update(update_data).eq("id", str(project_id)).execute()
    return result.data[0] if result.data else {}


# ============================================================
# FEEDBACK
# ============================================================

feedback_router = APIRouter(prefix="/feedback", tags=["feedback"])


@feedback_router.post("/paste", response_model=FeedbackUploadResponse)
async def paste_feedback(
    project_id: UUID,
    body: FeedbackPaste,
    user: Annotated[dict, Depends(get_current_user)],
):
    """Paste raw feedback text for a project."""
    db = get_supabase_admin()
    _verify_project_access(db, project_id, user)

    rows = [
        {
            "project_id": str(project_id),
            "raw_text": text.strip(),
            "metadata": body.metadata or {},
            "status": "raw",
        }
        for text in body.texts
        if text.strip()
    ]

    if rows:
        db.table("feedback_items").insert(rows).execute()

    return FeedbackUploadResponse(
        items_created=len(rows),
        items_skipped=len(body.texts) - len(rows),
        source_id=None,
        session_id=None,
    )


@feedback_router.post("/upload", response_model=FeedbackUploadResponse)
async def upload_feedback_file(
    project_id: UUID,
    user: Annotated[dict, Depends(get_current_user)],
    file: UploadFile = File(...),  # noqa: B008
):
    """Upload a file containing feedback (CSV, TXT, DOCX, PDF)."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Enforce file size limit
    content = await file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 10 MB)")

    db = get_supabase_admin()
    _verify_project_access(db, project_id, user)

    texts: list[str] = []

    # Parse based on file type
    filename = file.filename.lower()
    if filename.endswith(".txt"):
        texts = content.decode("utf-8", errors="ignore").split("\n\n")
    elif filename.endswith(".csv"):
        import csv
        import io
        reader = csv.reader(io.StringIO(content.decode("utf-8", errors="ignore")))
        next(reader, None)  # skip header row
        for row in reader:
            texts.append(" | ".join(row))
    elif filename.endswith(".docx"):
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
    elif filename.endswith(".pdf"):
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")

    # Batch insert
    rows = [
        {
            "project_id": str(project_id),
            "raw_text": text.strip(),
            "status": "raw",
            "metadata": {"filename": file.filename},
        }
        for text in texts
        if text.strip()
    ]

    if rows:
        db.table("feedback_items").insert(rows).execute()

    return FeedbackUploadResponse(
        items_created=len(rows),
        items_skipped=len(texts) - len(rows),
        source_id=None,
        session_id=None,
    )


@feedback_router.post("/parse-file", response_model=FileParseResponse)
async def parse_feedback_file(
    file: UploadFile = File(...),  # noqa: B008
):
    """Extract texts from a file without inserting into the database."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    content = await file.read()
    if len(content) > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 10 MB)")

    texts: list[str] = []
    filename = file.filename.lower()

    if filename.endswith(".txt") or filename.endswith(".md"):
        texts = content.decode("utf-8", errors="ignore").split("\n\n")
    elif filename.endswith(".csv"):
        import csv
        import io
        reader = csv.reader(io.StringIO(content.decode("utf-8", errors="ignore")))
        next(reader, None)
        for row in reader:
            texts.append(" | ".join(row))
    elif filename.endswith(".tsv"):
        import io
        lines = content.decode("utf-8", errors="ignore").split("\n")
        for line in lines[1:]:
            if line.strip():
                texts.append(" | ".join(line.split("\t")))
    elif filename.endswith(".docx"):
        import io
        from docx import Document
        doc = Document(io.BytesIO(content))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
    elif filename.endswith(".pdf"):
        import io
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
    elif filename.endswith(".json"):
        import json
        parsed = json.loads(content.decode("utf-8", errors="ignore"))
        items = parsed if isinstance(parsed, list) else [parsed]
        text_fields = [
            "feedback_text", "feedback", "text", "content", "message",
            "body", "comment", "review", "note", "description", "summary",
            "response", "input", "query", "question", "answer",
        ]
        for item in items:
            if isinstance(item, str):
                texts.append(item)
            elif isinstance(item, dict):
                for field in text_fields:
                    if field in item and isinstance(item[field], str) and item[field].strip():
                        texts.append(item[field].strip())
                        break
                else:
                    texts.append(json.dumps(item))
    elif filename.endswith(".jsonl"):
        import json
        text_fields = [
            "feedback_text", "feedback", "text", "content", "message",
            "body", "comment", "review", "note", "description", "summary",
            "response", "input", "query", "question", "answer",
        ]
        for line in content.decode("utf-8", errors="ignore").split("\n"):
            if not line.strip():
                continue
            obj = json.loads(line)
            if isinstance(obj, str):
                texts.append(obj)
            elif isinstance(obj, dict):
                for field in text_fields:
                    if field in obj and isinstance(obj[field], str) and obj[field].strip():
                        texts.append(obj[field].strip())
                        break
                else:
                    texts.append(json.dumps(obj))
    else:
        # Fallback: try to read as text
        try:
            texts = content.decode("utf-8").split("\n")
        except UnicodeDecodeError:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")

    return FileParseResponse(texts=[t.strip() for t in texts if t.strip()])


@feedback_router.get("", response_model=FeedbackListResponse)
async def list_feedback(
    project_id: UUID,
    user: Annotated[dict, Depends(get_current_user)],
    page: int = Query(default=1, ge=1),
    page_size: int = Query(default=50, ge=1, le=200),
    status_filter: str | None = None,
):
    """List feedback items for a project."""
    db = get_supabase_admin()
    _verify_project_access(db, project_id, user)

    query = db.table("feedback_items").select("*", count="exact").eq("project_id", str(project_id))

    if status_filter:
        query = query.eq("status", status_filter)

    offset = (page - 1) * page_size
    result = query.order("created_at", desc=True).range(offset, offset + page_size - 1).execute()

    return FeedbackListResponse(
        items=[FeedbackItemResponse(**item) for item in result.data],
        total=result.count or 0,
        page=page,
        page_size=page_size,
    )
